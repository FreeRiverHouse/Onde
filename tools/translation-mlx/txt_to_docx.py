
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Config
INPUT_FILE = "tools/translation-mlx/capussela_output/traduzione_finale.txt"
OUTPUT_FILE = "tools/translation-mlx/capussela_output/Capussela_Traduzione.docx"
MAX_HEADING_CHARS = 100

def set_font_size(paragraph, size):
    for run in paragraph.runs:
        run.font.size = Pt(size)

def is_title(text):
    text_lower = text.strip().lower()
    
    # Check 1: Length (Titles are short)
    if len(text) > 80:
        return False

    # Check 2: Explicit "Capitolo" or "Parte"
    if text_lower.startswith("capitolo") or text_lower.startswith("parte") or text_lower.startswith("sezione"):
        return True

    # Check 3: Standard Titles
    standard_titles = ["introduzione", "prefazione", "conclusione", "indice", "sommario", 
                       "bibliografia", "riferimenti", "note", "ringraziamenti", "abstract"]
    
    # Exact match or title with colon (e.g. "Introduzione: ...")
    if text_lower in standard_titles:
        return True
    
    if any(text_lower.startswith(t) for t in standard_titles) and len(text) < 50:
        return True

    # Check 4: Uppercase SHORT titles (e.g. "THE END")
    if text.isupper() and len(text) < 40 and len(text) > 3:
        return True

    return False

print(f"Generando documento Word da: {INPUT_FILE}")

document = Document()

# Stile Base
style = document.styles['Normal']
font = style.font
font.name = 'Georgia'  # Più "libro" di Calibri
font.size = Pt(11)

# Titolo Principale
document.add_heading('Traduzione Automatica - MLX Qwen3-32B', 0)
document.add_paragraph('Generata da Pipeline Top G v12.0').alignment = WD_ALIGN_PARAGRAPH.CENTER

with open(INPUT_FILE, "r", encoding="utf-8") as f:
    text = f.read()

paragraphs = text.split('\n\n')

for para_text in paragraphs:
    para_text = para_text.strip()
    if not para_text:
        continue

    # Clean up formatting artifacts
    para_text = para_text.replace('\n', ' ')
    
    # Check if Heading
    if is_title(para_text):
        h = document.add_heading(para_text, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p = document.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Interlinea 1.15
        p.paragraph_format.line_spacing = 1.15
        # Spazio dopo paragrafo
        p.paragraph_format.space_after = Pt(6)

print(f"Saving to {OUTPUT_FILE}...")
document.save(OUTPUT_FILE)
print("Fatto! Documento pronto.")
