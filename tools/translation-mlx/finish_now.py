
import sys
import os
import subprocess

# PORTS
BASE_DIR = "tools/translation-mlx/capussela_output"
ORIGINAL_FILE = os.path.join(BASE_DIR, "traduzione_finale.txt")
V5_FILE = os.path.join(BASE_DIR, "traduzione_FINAL_V5.txt")
MERGED_TXT = os.path.join(BASE_DIR, "traduzione_COMPLETA_FINAL.txt")
OUTPUT_DOCX = os.path.join(BASE_DIR, "Capussela_Traduzione_FINAL.docx")

print(f"--- STARTING FINAL ASSEMBLY ---")

# 1. MERGE
try:
    with open(ORIGINAL_FILE, "r", encoding="utf-8") as f:
        orig_paras = f.read().split('\n\n')
except FileNotFoundError:
    print(f"CRITICAL: Original file not found at {ORIGINAL_FILE}")
    orig_paras = []

try:
    with open(V5_FILE, "r", encoding="utf-8") as f:
        v5_paras = f.read().split('\n\n')
    print(f"V5 Paras Found: {len(v5_paras)}")
except FileNotFoundError:
    print(f"WARNING: V5 file not found. Using ONLY Original.")
    v5_paras = []

# Logic
final_paras = list(v5_paras)
if len(orig_paras) > len(v5_paras):
    missing = len(orig_paras) - len(v5_paras)
    print(f"Appending {missing} paragraphs from Original...")
    final_paras.extend(orig_paras[len(v5_paras):])
else:
    print("V5 covers the whole book.")

# Save Merged
with open(MERGED_TXT, "w", encoding="utf-8") as f:
    f.write('\n\n'.join(final_paras))
print(f"Merged Text Saved: {MERGED_TXT}")

# 2. GENERATE DOCX (Try python-docx first)
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    print("Generating DOCX with python-docx...")
    doc = Document()
    
    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    
    doc.add_heading('Capussela - The Republic of Innovation (IT)', 0)
    
    for i, p_text in enumerate(final_paras):
        p_text = p_text.strip()
        if not p_text: continue
        
        # Simple Title Heuristic
        is_title = False
        if len(p_text) < 100 and not p_text.endswith('.'):
             if any(x in p_text.lower() for x in ['capitolo', 'chapter', 'introduzione', 'prefazione', 'indice', 'bibliografia']):
                 is_title = True
             elif p_text.isupper() and len(p_text) > 3:
                 is_title = True
        
        if is_title:
            doc.add_heading(p_text, level=1)
        else:
            p = doc.add_paragraph(p_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            
    doc.save(OUTPUT_DOCX)
    print(f"SUCCESS: DOCX generated at {OUTPUT_DOCX}")

except ImportError:
    print("python-docx not found. Falling back to native textutil...")
    # Fallback
    subprocess.run(["textutil", "-convert", "docx", "-output", OUTPUT_DOCX, MERGED_TXT])
    print(f"SUCCESS (Fallback): DOCX generated at {OUTPUT_DOCX}")

except Exception as e:
    print(f"ERROR with python-docx: {e}")
    print("Falling back to native textutil...")
    subprocess.run(["textutil", "-convert", "docx", "-output", OUTPUT_DOCX, MERGED_TXT])
    print(f"SUCCESS (Fallback): DOCX generated at {OUTPUT_DOCX}")
